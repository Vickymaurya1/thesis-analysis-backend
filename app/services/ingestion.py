import re
import json
import random
import os
import httpx
import pdfplumber
import docx
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from anthropic import Anthropic

from app.config import settings
from app.models import ThesisVersion, Chunk, CorpusDocument, CorpusSourceEnum
from app.services.llm import llm_service

# Define SECTION synonyms pattern
SECTION_SYNONYMS = {
    "introduction": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Introduction|Motivation)\b', re.IGNORECASE),
    "related_work": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Related Work|Literature Review|Background)\b', re.IGNORECASE),
    "methodology": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Methodology|Research Methodology|Methods|System Design|Experimental Setup)\b', re.IGNORECASE),
    "results": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Results|Evaluation|Experimental Results|Findings|Experimental Evaluation)\b', re.IGNORECASE),
    "discussion": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Discussion|Analysis)\b', re.IGNORECASE),
    "conclusion": re.compile(r'^(?:Chapter\s+\d+\s+)?(?:Conclusion|Conclusion and Future Work|Summary)\b', re.IGNORECASE)
}

def extract_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        text_pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
        return "\n\n".join(text_pages)
    elif ext == '.docx':
        doc = docx.Document(file_path)
        paragraphs_text = [p.text for p in doc.paragraphs]
        return "\n\n".join(paragraphs_text)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def detect_sections_regex(raw_text: str) -> Dict[str, List[int]]:
    lines = raw_text.split('\n')
    offsets = []
    current_offset = 0
    for line in lines:
        offsets.append((current_offset, line))
        current_offset += len(line) + 1  # count newline
        
    section_starts = {}
    for offset, line in offsets:
        clean_line = line.strip()
        if not clean_line or len(clean_line) > 120:
            continue
            
        # Strip numbering or headers prefix
        normalized_line = re.sub(r'^(?:Chapter\s+\d+|\d+(?:\.\d+)*|#+)\s*', '', clean_line, flags=re.IGNORECASE).strip()
        
        for section_name, pattern in SECTION_SYNONYMS.items():
            if section_name in section_starts:
                continue
            if pattern.match(normalized_line):
                section_starts[section_name] = offset
                break
                
    if not section_starts:
        return {}
        
    # Convert offsets to ranges
    sorted_sections = sorted(section_starts.items(), key=lambda x: x[1])
    section_map = {}
    for idx, (sec_name, start_offset) in enumerate(sorted_sections):
        end_offset = len(raw_text)
        if idx + 1 < len(sorted_sections):
            end_offset = sorted_sections[idx + 1][1]
        section_map[sec_name] = [start_offset, end_offset]
        
    return section_map

def extract_candidate_headings(raw_text: str) -> List[Dict[str, Any]]:
    lines = raw_text.split('\n')
    candidates = []
    current_offset = 0
    for idx, line in enumerate(lines):
        clean_line = line.strip()
        line_len = len(line)
        is_candidate = False
        
        if clean_line and len(clean_line) < 120:
            # Check prefix triggers or uppercase or double line break boundaries
            prev_empty = (idx == 0 or not lines[idx-1].strip())
            next_empty = (idx == len(lines)-1 or not lines[idx+1].strip())
            starts_with_heading = bool(re.match(r'^(?:Chapter\s+\d+|\d+(?:\.\d+)*|#+)\s', clean_line, re.IGNORECASE))
            is_uppercase = clean_line.isupper()
            
            if starts_with_heading or is_uppercase or (prev_empty and next_empty):
                is_candidate = True
                
        if is_candidate:
            candidates.append({
                "offset": current_offset,
                "text": clean_line
            })
        current_offset += line_len + 1
    return candidates

def detect_sections_llm(raw_text: str) -> Dict[str, List[int]]:
    candidates = extract_candidate_headings(raw_text)
    if not candidates:
        return {}

    # If Anthropic API key is mock or absent, we stub LLM fallback for tests
    api_key = settings.ANTHROPIC_API_KEY
    if not api_key or api_key == "mock_api_key_for_testing" or settings.ENV == "test":
        # Returns simple fallback map based on candidate presence or heuristic
        mock_map = {}
        for c in candidates:
            text = c["text"].lower()
            for key in SECTION_SYNONYMS.keys():
                if key not in mock_map and key[:4] in text:
                    mock_map[key] = c["offset"]
        
        # Sort and construct ranges
        sorted_map = sorted(mock_map.items(), key=lambda x: x[1])
        res = {}
        for idx, (sec_name, start) in enumerate(sorted_map):
            end = len(raw_text)
            if idx + 1 < len(sorted_map):
                end = sorted_map[idx+1][1]
            res[sec_name] = [start, end]
        return res

    # Otherwise run proper Claude query
    system_prompt = (
        "You are a parser assistant for academic theses.\n"
        "Your task is to analyze a list of candidate headings from a thesis and map them to the following 6 core sections:\n"
        "1. introduction\n"
        "2. related_work\n"
        "3. methodology\n"
        "4. results\n"
        "5. discussion\n"
        "6. conclusion\n\n"
        "You will be given a JSON array of candidate headings, where each item has an 'offset' and 'text'.\n"
        "Select the entry that best represents the beginning of each core section. Return the result in a JSON object mapping the section name to its start offset.\n"
        "If a section is not found or not present in the candidates, do not include it in your output.\n\n"
        "Output strict JSON only, with no explanation or conversational text."
    )
    
    user_content = (
        f"Candidate headings:\n{json.dumps(candidates, indent=2)}\n\n"
        f"Return JSON mapping core section names to their start offsets, e.g.:\n"
        f"{{\"introduction\": 1045, \"related_work\": 12450, ...}}"
    )
    
    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-3-5-sonnet-20241022",
        max_tokens=2000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_content}],
        temperature=0.0
    )
    
    raw_res = response.content[0].text.strip()
    try:
        if "```json" in raw_res:
            raw_res = raw_res.split("```json")[1].split("```")[0].strip()
        elif "```" in raw_res:
            raw_res = raw_res.split("```")[1].split("```")[0].strip()
        offsets_map = json.loads(raw_res)
        
        # Resolve offset ranges
        sorted_offsets = sorted(offsets_map.items(), key=lambda x: x[1])
        section_map = {}
        for idx, (sec_name, start_offset) in enumerate(sorted_offsets):
            end_offset = len(raw_text)
            if idx + 1 < len(sorted_offsets):
                end_offset = sorted_offsets[idx + 1][1]
            section_map[sec_name] = [start_offset, end_offset]
        return section_map
    except Exception as e:
        raise ValueError(f"Failed to parse LLM section map output: {e}. Raw: {raw_res}")

def get_embeddings(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
        
    api_key = settings.VOYAGE_API_KEY
    if not api_key or api_key == "mock_api_key_for_testing" or settings.ENV == "test":
        # Refuse to start/fail loudly outside test mode if credentials are missing
        if settings.ENV != "test":
            raise ValueError("VOYAGE_API_KEY is missing/mock in a non-test environment.")
        # Test mock stub fallback
        return [[random.random() for _ in range(1024)] for _ in texts]
        
    url = "https://api.voyageai.com/v1/embeddings"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Voyage AI supports batch embeddings
    body = {
        "input": texts,
        "model": "voyage-3"
    }
    
    with httpx.Client() as client:
        response = client.post(url, headers=headers, json=body, timeout=20.0)
        if response.status_code == 200:
            data = response.json()
            return [item["embedding"] for item in data["data"]]
        else:
            raise ValueError(f"Voyage API call failed: {response.status_code} - {response.text}")

def split_into_sentences(text: str) -> List[str]:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def chunk_section_text(section_text: str, section_label: str) -> List[Dict[str, Any]]:
    paragraphs = section_text.split('\n\n')
    chunks = []
    current_chunk = []
    current_len = 0
    chunk_idx = 0
    
    for para in paragraphs:
        para_text = para.strip()
        if not para_text:
            continue
            
        para_len = len(para_text)
        
        # If single paragraph is larger than chunk limit, split it by sentences
        if para_len > 2500:
            if current_chunk:
                chunks.append({
                    "chunk_index": chunk_idx,
                    "content": "\n\n".join(current_chunk),
                    "section_label": section_label
                })
                chunk_idx += 1
                current_chunk = []
                current_len = 0
                
            sentences = split_into_sentences(para_text)
            temp_chunk = []
            temp_len = 0
            for sent in sentences:
                if temp_len + len(sent) > 2000:
                    chunks.append({
                        "chunk_index": chunk_idx,
                        "content": " ".join(temp_chunk),
                        "section_label": section_label
                    })
                    chunk_idx += 1
                    temp_chunk = []
                    temp_len = 0
                temp_chunk.append(sent)
                temp_len += len(sent) + 1
            if temp_chunk:
                chunks.append({
                    "chunk_index": chunk_idx,
                    "content": " ".join(temp_chunk),
                    "section_label": section_label
                })
                chunk_idx += 1
            continue
            
        if current_len + para_len > 2000:
            chunks.append({
                "chunk_index": chunk_idx,
                "content": "\n\n".join(current_chunk),
                "section_label": section_label
            })
            chunk_idx += 1
            current_chunk = []
            current_len = 0
            
        current_chunk.append(para_text)
        current_len += para_len + 2
        
    if current_chunk:
        chunks.append({
            "chunk_index": chunk_idx,
            "content": "\n\n".join(current_chunk),
            "section_label": section_label
        })
        
    return chunks

def compute_paragraph_diff(raw_text1: str, map1: Dict[str, List[int]], raw_text2: str, map2: Dict[str, List[int]]) -> Dict[str, Any]:
    changed_paragraphs = []
    changed_sections = []
    
    # We compare paragraph-by-paragraph for each section in both section maps
    for sec_name in map2.keys():
        if sec_name not in map1:
            changed_sections.append(sec_name)
            continue
            
        start1, end1 = map1[sec_name]
        start2, end2 = map2[sec_name]
        
        text1 = raw_text1[start1:end1].strip()
        text2 = raw_text2[start2:end2].strip()
        
        paras1 = [p.strip() for p in text1.split('\n\n') if p.strip()]
        paras2 = [p.strip() for p in text2.split('\n\n') if p.strip()]
        
        sec_changed = False
        for idx, p2 in enumerate(paras2):
            # Check if paragraph at index is different, or if index is out of bounds for paras1
            if idx >= len(paras1) or paras1[idx] != p2:
                changed_paragraphs.append({
                    "index": idx,
                    "section": sec_name,
                    "text": p2
                })
                sec_changed = True
                
        if sec_changed:
            changed_sections.append(sec_name)
            
    return {
        "changed_paragraphs": changed_paragraphs,
        "changed_sections": list(set(changed_sections))
    }

async def run_ingestion_pipeline(db: Session, version: ThesisVersion, file_path: str):
    # 1. Extract text
    raw_text = extract_text_from_file(file_path)
    version.raw_text = raw_text
    
    # 2. Section detection (Synonyms Regex first, LLM fallback)
    section_map = detect_sections_regex(raw_text)
    if not section_map:
        section_map = detect_sections_llm(raw_text)
        
    version.section_map = section_map
    
    # 3. Chunking sections
    all_chunks = []
    for sec_name, range_vals in section_map.items():
        start, end = range_vals
        section_text = raw_text[start:end]
        section_chunks = chunk_section_text(section_text, sec_name)
        all_chunks.extend(section_chunks)
        
    # 4. Generate Embeddings (batch call)
    chunk_contents = [c["content"] for c in all_chunks]
    embeddings = get_embeddings(chunk_contents)
    
    # 5. Create chunk rows and mirror to CorpusDocument
    for idx, c in enumerate(all_chunks):
        embedding_vector = embeddings[idx] if idx < len(embeddings) else None
        db_chunk = Chunk(
            version_id=version.id,
            chunk_index=c["chunk_index"],
            content=c["content"],
            section_label=c["section_label"],
            embedding=embedding_vector
        )
        db.add(db_chunk)
        
        corp_doc = CorpusDocument(
            source_type=CorpusSourceEnum.internal_thesis,
            source_thesis_id=version.thesis_id,
            title=version.thesis.title if version.thesis else None,
            chunk_text=c["content"],
            embedding=embedding_vector
        )
        db.add(corp_doc)
        
    # 6. Compute diff if version > 1
    if version.version_number > 1:
        prev_version = db.query(ThesisVersion).filter(
            ThesisVersion.thesis_id == version.thesis_id,
            ThesisVersion.version_number == version.version_number - 1
        ).first()
        if prev_version and prev_version.section_map and prev_version.raw_text:
            diff = compute_paragraph_diff(
                prev_version.raw_text, prev_version.section_map,
                version.raw_text, version.section_map
            )
            version.diff_from_prev = diff
            
    db.commit()
