import pytest
import datetime
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.models import (
    User, Thesis, ThesisVersion, CorpusDocument, CorpusSourceEnum,
    Flag, FlagTypeEnum, AnalysisSnapshot, SemanticScholarSearch, RoleEnum, SeverityEnum
)
from app.services.ingestion import run_ingestion_pipeline
from app.services.plagiarism import run_plagiarism_review_pipeline
from app.services.novelty import run_novelty_review_pipeline

@pytest.mark.asyncio
async def test_corpus_mirroring_and_no_self_matching(client: TestClient, db: Session):
    # Setup student
    stud = User(email="stud_corp@uni.edu", password_hash="hash", role=RoleEnum.student, name="Stud")
    db.add(stud)
    db.commit()

    # Create thesis
    thesis = Thesis(owner_id=stud.id, title="Corpus Thesis", field="CS")
    db.add(thesis)
    db.commit()

    # Create mock PDF file
    import os
    os.makedirs("uploads", exist_ok=True)
    temp_path = "uploads/temp_test_mirror.docx"
    
    # Write mock docx using python-docx
    import docx
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("We propose a novel attention head pruning mechanism that optimizes transformer inference speed.")
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph("We evaluate the model on standard translation test sets.")
    doc.save(temp_path)

    # Register version
    version = ThesisVersion(thesis_id=thesis.id, version_number=1, file_ref=temp_path)
    db.add(version)
    db.commit()

    # Ingestion
    await run_ingestion_pipeline(db, version, temp_path)
    
    # Clean up temp file
    if os.path.exists(temp_path):
        os.remove(temp_path)

    # 1. Verify chunks are mirrored in CorpusDocument
    corp_docs = db.query(CorpusDocument).filter(CorpusDocument.source_thesis_id == thesis.id).all()
    assert len(corp_docs) > 0
    assert corp_docs[0].source_type == CorpusSourceEnum.internal_thesis
    assert corp_docs[0].title == "Corpus Thesis"

    # 2. Run plagiarism pipeline
    snapshot = run_plagiarism_review_pipeline(db, version)
    
    # 3. Assert self-exclusion (no self matches found)
    # The flag count should be 0 because we filter out CorpusDocuments from this same thesis!
    assert snapshot.scores["flag_count"] == 0
    assert len(snapshot.scores["verdicts"]) == 0

@pytest.mark.asyncio
async def test_plagiarism_verdict_threshold_flagging(db: Session):
    stud1 = User(email="std1@uni.edu", password_hash="hash", role=RoleEnum.student, name="Stud1")
    stud2 = User(email="std2@uni.edu", password_hash="hash", role=RoleEnum.student, name="Stud2")
    db.add_all([stud1, stud2])
    db.commit()

    thesis1 = Thesis(owner_id=stud1.id, title="Original Work")
    thesis2 = Thesis(owner_id=stud2.id, title="Plagiarized Work")
    db.add_all([thesis1, thesis2])
    db.commit()

    # Seed peer student work in CorpusDocument
    peer_vector = [0.1] * 1024
    peer_doc = CorpusDocument(
        source_type=CorpusSourceEnum.internal_thesis,
        source_thesis_id=thesis1.id,
        title="Original Work",
        chunk_text="We split our dataset into 80% training and 20% testing.",
        embedding=peer_vector
    )
    db.add(peer_doc)
    db.commit()

    # Chunk 1 (Coincidental overlap trigger): "split dataset 80/20"
    v2 = ThesisVersion(thesis_id=thesis2.id, version_number=1, raw_text="We split our dataset 80% training.")
    db.add(v2)
    db.commit()
    db.refresh(v2)

    from app.models import Chunk
    # Insert chunk matching peer vector
    chunk1 = Chunk(
        version_id=v2.id,
        chunk_index=0,
        content="We split our dataset 80% training.",
        section_label="methodology",
        embedding=peer_vector
    )
    db.add(chunk1)
    db.commit()

    db.refresh(v2)
    snapshot = run_plagiarism_review_pipeline(db, v2)
    
    # Verify that coincidental_overlap verdict is generated but NO FLAG is created
    assert len(snapshot.scores["verdicts"]) == 1
    assert snapshot.scores["verdicts"][0]["verdict"] == "coincidental_overlap"
    assert snapshot.scores["flag_count"] == 0
    
    flags = db.query(Flag).filter(Flag.snapshot_id == snapshot.id).all()
    assert len(flags) == 0

    # Let's write a version that has a direct copy trigger (doesn't contain coincidental keywords)
    v3 = ThesisVersion(thesis_id=thesis2.id, version_number=2, raw_text="We copied this text exactly.")
    db.add(v3)
    db.commit()
    
    chunk2 = Chunk(
        version_id=v3.id,
        chunk_index=0,
        content="We copied this text exactly.",
        section_label="introduction",
        embedding=peer_vector
    )
    db.add(chunk2)
    db.commit()

    db.refresh(v3)
    snapshot2 = run_plagiarism_review_pipeline(db, v3)
    
    # Verdict likely_copied -> should generate a Flag!
    assert snapshot2.scores["flag_count"] == 1
    assert snapshot2.scores["verdicts"][0]["verdict"] == "likely_copied"
    
    flags2 = db.query(Flag).filter(Flag.snapshot_id == snapshot2.id).all()
    assert len(flags2) == 1
    assert flags2[0].type == FlagTypeEnum.plagiarism

def test_student_privacy_masking(client: TestClient, db: Session):
    # Register student
    s_res = client.post("/auth/register", json={
        "email": "student_privacy@uni.edu",
        "password": "pass",
        "role": "student",
        "name": "Privacy Stud",
        "university": "Uni",
        "department": "CS",
        "degree": "BTech",
        "field_of_study": "CS"
    })
    student_id = s_res.json()["id"]

    # Register teacher
    t_res = client.post("/auth/register", json={
        "email": "teacher_privacy@uni.edu",
        "password": "pass",
        "role": "teacher",
        "name": "Privacy Prof",
        "university": "Uni",
        "department": "CS",
        "designation": "Professor"
    })
    teacher_id = t_res.json()["id"]

    # Log in as Student
    s_login = client.post("/auth/login", data={"username": "student_privacy@uni.edu", "password": "pass"})
    s_token = s_login.json()["access_token"]
    s_headers = {"Authorization": f"Bearer {s_token}"}

    # Log in as Teacher
    t_login = client.post("/auth/login", data={"username": "teacher_privacy@uni.edu", "password": "pass"})
    t_token = t_login.json()["access_token"]
    t_headers = {"Authorization": f"Bearer {t_token}"}

    # Create student's thesis and associate advisor
    thesis = Thesis(owner_id=student_id, advisor_id=teacher_id, title="Privacy Thesis", field="CS")
    db.add(thesis)
    db.commit()

    version = ThesisVersion(thesis_id=thesis.id, version_number=1, raw_text="Sample raw text.")
    db.add(version)
    db.commit()
    
    thesis.current_version_id = version.id
    db.commit()

    # Create dummy plagiarism flag matching internal thesis
    snapshot = AnalysisSnapshot(version_id=version.id, tool_type="plagiarism", scores={})
    db.add(snapshot)
    db.commit()

    flag = Flag(
        snapshot_id=snapshot.id,
        type=FlagTypeEnum.plagiarism,
        severity=SeverityEnum.critical,
        message="Matches internal thesis [ID: peer_thesis_id_123, Title: 'Peer Student Hidden Work'] | Copy detected.",
        evidence_excerpt="Sample copy"
    )
    db.add(flag)
    db.commit()

    # Retrieve flags as Student -> details must be masked
    s_flags = client.get(f"/theses/{thesis.id}/flags", headers=s_headers).json()
    assert len(s_flags) == 1
    assert "Details restricted to advisor" in s_flags[0]["message"]
    assert "Peer Student Hidden Work" not in s_flags[0]["message"]

    # Retrieve flags as Teacher -> details must NOT be masked
    t_flags = client.get(f"/theses/{thesis.id}/flags", headers=t_headers).json()
    assert len(t_flags) == 1
    assert "Peer Student Hidden Work" in t_flags[0]["message"]
    assert "Details restricted to advisor" not in t_flags[0]["message"]

@pytest.mark.asyncio
async def test_novelty_detection_and_query_caching(client: TestClient, db: Session):
    student = User(email="novelty_stud@uni.edu", password_hash="hash", role=RoleEnum.student, name="Stud")
    db.add(student)
    db.commit()
    
    thesis = Thesis(
        owner_id=student.id,
        title="Pruning Transformers at Scale",
        field="CS",
        degree_level="BTech"
    )
    db.add(thesis)
    db.commit()
    
    # Contribution claims in introduction section
    version = ThesisVersion(
        thesis_id=thesis.id,
        version_number=1,
        raw_text="Introduction\nWe propose a novel attention head pruning mechanism.",
        section_map={"introduction": [0, 60]}
    )
    db.add(version)
    db.commit()
    
    thesis.current_version_id = version.id
    db.commit()

    from app.models import Chunk
    chunk = Chunk(
        version_id=version.id,
        chunk_index=0,
        content="We propose a novel attention head pruning mechanism.",
        section_label="introduction",
        embedding=[0.1]*1024
    )
    db.add(chunk)
    db.commit()

    # 1. Run novelty pipeline
    snapshot = run_novelty_review_pipeline(db, version)
    
    # 2. Check that SemanticScholarSearch query was cached
    search_log = db.query(SemanticScholarSearch).filter(
        SemanticScholarSearch.query_text == "cs pruning transformers at scale"
    ).first()
    assert search_log is not None
    assert search_log.fetched_at is not None
    
    # Remember the fetch timestamp
    first_fetch = search_log.fetched_at

    # 3. Running novelty pipeline again should read from query cache and not query/update fetched_at
    snapshot_2 = run_novelty_review_pipeline(db, version)
    db.refresh(search_log)
    assert search_log.fetched_at == first_fetch  # Timestamp did not change because search was cached!
    
    # 4. Verify novelty flagging rules: "pruning" contribution yields "substantially_overlapping" and generates Flag
    assert len(snapshot.scores["claims_reviewed"]) == 1
    assert snapshot.scores["claims_reviewed"][0]["verdict"] == "substantially_overlapping"
    
    flags = db.query(Flag).filter(Flag.snapshot_id == snapshot.id).all()
    assert len(flags) == 1
    assert flags[0].type == FlagTypeEnum.novelty
    assert flags[0].severity == SeverityEnum.critical
